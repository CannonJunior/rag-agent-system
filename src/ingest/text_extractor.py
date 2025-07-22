import asyncio
import hashlib
import mimetypes
from pathlib import Path
from typing import Optional, Dict, Any, List, Tuple
from dataclasses import dataclass

import aiofiles
from loguru import logger
from pypdf import PdfReader
from docx import Document
import textract

from config.settings import settings


@dataclass
class ExtractedContent:
    """Container for extracted text content."""
    text: str
    metadata: Dict[str, Any]
    content_hash: str
    file_size: int
    content_type: str
    extraction_method: str
    page_count: Optional[int] = None
    word_count: Optional[int] = None
    char_count: Optional[int] = None


class TextExtractor:
    """Extract text content from various file formats."""
    
    def __init__(self):
        self.supported_types = {
            '.txt': self._extract_text,
            '.md': self._extract_text, 
            '.rst': self._extract_text,
            '.py': self._extract_text,
            '.js': self._extract_text,
            '.ts': self._extract_text,
            '.json': self._extract_text,
            '.yaml': self._extract_text,
            '.yml': self._extract_text,
            '.xml': self._extract_text,
            '.html': self._extract_text,
            '.csv': self._extract_text,
            '.pdf': self._extract_pdf,
            '.docx': self._extract_docx,
            '.doc': self._extract_with_textract,
            '.rtf': self._extract_with_textract,
            '.odt': self._extract_with_textract,
        }
    
    async def extract(self, file_path: Path) -> Optional[ExtractedContent]:
        """Extract text from a file."""
        try:
            if not file_path.exists():
                logger.error(f"File not found: {file_path}")
                return None
            
            # Check file size
            file_size = file_path.stat().st_size
            max_size = settings.ingest.max_file_size_mb * 1024 * 1024
            if file_size > max_size:
                logger.warning(f"File too large ({file_size} bytes): {file_path}")
                return None
            
            # Check file extension
            extension = file_path.suffix.lower()
            if extension not in settings.ingest.supported_extensions:
                logger.warning(f"Unsupported file type {extension}: {file_path}")
                return None
            
            # Get content type
            content_type, _ = mimetypes.guess_type(str(file_path))
            if not content_type:
                content_type = f"application/{extension[1:]}" if extension else "application/octet-stream"
            
            # Extract text using appropriate method
            extractor = self.supported_types.get(extension)
            if not extractor:
                logger.warning(f"No extractor for {extension}: {file_path}")
                return None
            
            logger.info(f"Extracting text from {file_path} using {extractor.__name__}")
            
            text, metadata = await extractor(file_path)
            
            if not text or not text.strip():
                logger.warning(f"No text extracted from {file_path}")
                return None
            
            # Generate content hash
            content_hash = hashlib.sha256(text.encode('utf-8')).hexdigest()
            
            # Calculate metrics
            word_count = len(text.split())
            char_count = len(text)
            
            return ExtractedContent(
                text=text,
                metadata=metadata,
                content_hash=content_hash,
                file_size=file_size,
                content_type=content_type,
                extraction_method=extractor.__name__,
                word_count=word_count,
                char_count=char_count,
                page_count=metadata.get('page_count')
            )
            
        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {e}")
            return None
    
    async def _extract_text(self, file_path: Path) -> Tuple[str, Dict[str, Any]]:
        """Extract text from plain text files."""
        try:
            async with aiofiles.open(file_path, 'r', encoding='utf-8') as f:
                text = await f.read()
            
            metadata = {
                'encoding': 'utf-8',
                'line_count': len(text.splitlines())
            }
            
            return text, metadata
            
        except UnicodeDecodeError:
            # Try with different encodings
            for encoding in ['latin-1', 'cp1252', 'iso-8859-1']:
                try:
                    async with aiofiles.open(file_path, 'r', encoding=encoding) as f:
                        text = await f.read()
                    
                    metadata = {
                        'encoding': encoding,
                        'line_count': len(text.splitlines())
                    }
                    
                    return text, metadata
                except Exception:
                    continue
            
            raise ValueError(f"Could not decode text file with any encoding: {file_path}")
    
    async def _extract_pdf(self, file_path: Path) -> Tuple[str, Dict[str, Any]]:
        """Extract text from PDF files."""
        def _extract_pdf_sync():
            reader = PdfReader(str(file_path))
            
            # Extract metadata
            info = reader.metadata or {}
            metadata = {
                'page_count': len(reader.pages),
                'title': info.get('/Title', ''),
                'author': info.get('/Author', ''),
                'subject': info.get('/Subject', ''),
                'creator': info.get('/Creator', ''),
                'producer': info.get('/Producer', ''),
                'creation_date': str(info.get('/CreationDate', '')),
                'modification_date': str(info.get('/ModDate', ''))
            }
            
            # Extract text from all pages
            text_parts = []
            for page_num, page in enumerate(reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text.strip():
                        text_parts.append(f"--- Page {page_num + 1} ---\n{page_text}\n")
                except Exception as e:
                    logger.warning(f"Error extracting page {page_num + 1} from {file_path}: {e}")
                    continue
            
            text = '\n'.join(text_parts)
            return text, metadata
        
        # Run PDF extraction in thread pool to avoid blocking
        text, metadata = await asyncio.to_thread(_extract_pdf_sync)
        return text, metadata
    
    async def _extract_docx(self, file_path: Path) -> Tuple[str, Dict[str, Any]]:
        """Extract text from DOCX files."""
        def _extract_docx_sync():
            doc = Document(str(file_path))
            
            # Extract text from paragraphs
            text_parts = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(' | '.join(row_text))
            
            text = '\n'.join(text_parts)
            
            # Extract metadata
            core_props = doc.core_properties
            metadata = {
                'title': core_props.title or '',
                'author': core_props.author or '',
                'subject': core_props.subject or '',
                'keywords': core_props.keywords or '',
                'created': str(core_props.created) if core_props.created else '',
                'modified': str(core_props.modified) if core_props.modified else '',
                'paragraph_count': len(doc.paragraphs),
                'table_count': len(doc.tables)
            }
            
            return text, metadata
        
        text, metadata = await asyncio.to_thread(_extract_docx_sync)
        return text, metadata
    
    async def _extract_with_textract(self, file_path: Path) -> Tuple[str, Dict[str, Any]]:
        """Extract text using textract (fallback for various formats)."""
        def _extract_textract_sync():
            text = textract.process(str(file_path)).decode('utf-8')
            metadata = {
                'extraction_library': 'textract',
                'note': 'Advanced extraction method used'
            }
            return text, metadata
        
        try:
            text, metadata = await asyncio.to_thread(_extract_textract_sync)
            return text, metadata
        except Exception as e:
            logger.error(f"Textract extraction failed for {file_path}: {e}")
            raise
    
    def is_supported(self, file_path: Path) -> bool:
        """Check if file type is supported for extraction."""
        extension = file_path.suffix.lower()
        return extension in settings.ingest.supported_extensions
    
    async def validate_content(self, content: ExtractedContent) -> bool:
        """Validate extracted content quality."""
        if not content.text or len(content.text.strip()) < 10:
            return False
        
        # Check for reasonable text-to-size ratio
        if content.file_size > 0:
            ratio = len(content.text) / content.file_size
            if ratio < 0.001:  # Very low text extraction ratio
                logger.warning(f"Low text extraction ratio: {ratio}")
        
        # Check for mostly non-printable characters
        printable_chars = sum(1 for c in content.text if c.isprintable())
        printable_ratio = printable_chars / len(content.text)
        if printable_ratio < 0.8:
            logger.warning(f"Low printable character ratio: {printable_ratio}")
            return False
        
        return True


# Global text extractor instance
text_extractor = TextExtractor()